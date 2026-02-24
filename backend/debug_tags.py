import docx
import re
import os

def extract_tags_from_docx(path):
    if not os.path.exists(path):
        return set()
    doc = docx.Document(path)
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    full = " ".join(text)
    # Match {{ tag }} or {{tag}}
    return set(re.findall(r'\{\{\s*(.*?)\s*\}\}', full))

def extract_tags_from_md(path):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    return set(re.findall(r'\{\{\s*(.*?)\s*\}\}', content))

base_dir = os.path.dirname(__file__)
template_pl = os.path.join(base_dir, '..', 'template_pl.docx')
template_en = os.path.join(base_dir, '..', 'template_en.docx')
gemini_md = os.path.join(base_dir, '..', 'GEMINI.md')

tags_pl = extract_tags_from_docx(template_pl)
tags_en = extract_tags_from_docx(template_en)
tags_doc = extract_tags_from_md(gemini_md)

print("--- TAGS IN template_pl.docx ---")
print(sorted(list(tags_pl)))
print(f"Total: {len(tags_pl)}")

print("\n--- TAGS IN template_en.docx ---")
print(sorted(list(tags_en)))
print(f"Total: {len(tags_en)}")

print("\n--- TAGS DOCUMENTED IN GEMINI.md ---")
print(sorted(list(tags_doc)))
print(f"Total: {len(tags_doc)}")

common = tags_pl.intersection(tags_en)
only_pl = tags_pl - tags_en
only_en = tags_en - tags_pl

print("\n--- DISCREPANCIES BETWEEN TEMPLATES ---")
print(f"Only in PL: {only_pl}")
print(f"Only in EN: {only_en}")

undocumented_pl = tags_pl - tags_doc
unimplemented_pl = tags_doc - tags_pl

print("\n--- DISCREPANCIES (template_pl.docx vs GEMINI.md) ---")
print(f"In Template but NOT in GEMINI.md: {undocumented_pl}")
print(f"In GEMINI.md but NOT in Template: {unimplemented_pl}")

undocumented_en = tags_en - tags_doc
unimplemented_en = tags_doc - tags_en

print("\n--- DISCREPANCIES (template_en.docx vs GEMINI.md) ---")
print(f"In Template but NOT in GEMINI.md: {undocumented_en}")
print(f"In GEMINI.md but NOT in Template: {unimplemented_en}")
