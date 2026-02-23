import docx
import pdfplumber

def parse_docx(file_path: str):
    """Parses a .docx file and extracts text and tables."""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract tables and add their text to full_text for better parsing
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                    # Add cell text to full_text so extractors can see it
                    full_text.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
            
        return {
            "filename": file_path,
            "content": '\n'.join(full_text),
            "tables": tables
        }
    except Exception as e:
        return {"error": str(e)}

def parse_pdf(file_path: str):
    """Parses a .pdf file and extracts text and tables."""
    try:
        with pdfplumber.open(file_path) as pdf:
            full_text = []
            tables = []
            pages_data = []
            for page in pdf.pages:
                text = page.extract_text()
                page_tables = page.extract_tables()
                
                if text:
                    full_text.append(text)
                if page_tables:
                    tables.extend(page_tables)
                
                pages_data.append({
                    "text": text or "",
                    "tables": page_tables or []
                })

            return {
                "filename": file_path,
                "content": '\n'.join(full_text),
                "tables": tables,
                "pages": pages_data
            }
    except Exception as e:
        return {"error": str(e)}
