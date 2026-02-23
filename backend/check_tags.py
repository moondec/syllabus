import docx
import re
import os

try:
    doc_path = os.path.join(os.path.dirname(__file__), '..', 'template.docx')
    doc = docx.Document(doc_path)
    text = []
    
    for p in doc.paragraphs:
        text.append(p.text)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
                
    full = " ".join(text)
    tags = set(re.findall(r'\{\{\s*(.*?)\s*\}\}', full))
    
    print("Found Tags:", tags)
except Exception as e:
    print(f"Error reading docx: {e}")
