import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in kwargs[edge].items():
                element.set(qn('w:{}'.format(key)), str(val))

def rebuild():
    doc = Document()
    
    # Global Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SYLABUS (KARTA PRZEDMIOTU)')
    run.bold = True
    run.font.size = Pt(14)

    # Main Table
    # We'll use 1 column for now and add rows carefully, or a grid of 7 cols like identified before
    # Reference image looks like a mix of merged cells 
    # Let's use 7 columns as base grid
    table = doc.add_table(rows=0, cols=7)
    table.style = 'Table Grid'

    def add_full_row(text, bold=True, italic=False, font_size=10):
        row = table.add_row()
        cell = row.cells[0]
        # Merge all cells in row
        for i in range(1, 7):
            cell.merge(row.cells[i])
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        return cell

    def add_label_value_row(label, tag, label_cols=5):
        row = table.add_row()
        # Label cell
        label_cell = row.cells[0]
        for i in range(1, label_cols):
            label_cell.merge(row.cells[i])
        label_cell.text = label
        label_cell.paragraphs[0].runs[0].bold = True
        
        # Value cell
        value_cell = row.cells[label_cols]
        for i in range(label_cols + 1, 7):
            value_cell.merge(row.cells[i])
        value_cell.text = tag
        return row

    # R0: Nazwa przedmiotu + ECTS
    row0 = table.add_row()
    c0 = row0.cells[0]
    for i in range(1, 5): c0.merge(row0.cells[i])
    c0.text = "Nazwa przedmiotu (zgodna z zatwierdzonym programem studiów dla kierunku)"
    c0.paragraphs[0].runs[0].bold = True
    
    c5 = row0.cells[5]
    c5.merge(row0.cells[6])
    c5.text = "Liczba punktów ECTS"
    c5.paragraphs[0].runs[0].bold = True

    # R1: Value Row for Name and ECTS
    row1 = table.add_row()
    c0 = row1.cells[0]
    for i in range(1, 5): c0.merge(row1.cells[i])
    c0.text = "{{ name }}"
    
    c5 = row1.cells[5]
    c5.merge(row1.cells[6])
    c5.text = "{{ ects }}"

    # R2: Nazwa angielska label
    add_full_row("Nazwa przedmiotu w j. angielskim")
    # R3: Nazwa angielska value
    add_full_row("{{ name_en }}", bold=False)

    # R4: Jednostka label
    add_full_row("Jednostka(i) realizująca(e) przedmiot")
    # R5: Jednostka value
    add_full_row("{{ unit }}", bold=False)

    # R6: Kierownik label
    add_full_row("Kierownik przedmiotu")
    # R7: Kierownik value
    add_full_row("{{ kierownik }}", bold=False)

    # R8: Multi-column metadata Labels
    row8 = table.add_row()
    cells = row8.cells
    # Kierunek (2 cols)
    cells[0].merge(cells[1])
    cells[0].text = "Kierunek studiów"
    cells[0].paragraphs[0].runs[0].bold = True
    # Poziom (2 cols)
    cells[2].merge(cells[3])
    cells[2].text = "Poziom"
    cells[2].paragraphs[0].runs[0].bold = True
    # Profil (1 col)
    cells[4].text = "Profil"
    cells[4].paragraphs[0].runs[0].bold = True
    # Semestr (2 cols)
    cells[5].merge(cells[6])
    cells[5].text = "Semestr"
    cells[5].paragraphs[0].runs[0].bold = True

    # R9: Multi-column metadata Values
    row9 = table.add_row()
    cells = row9.cells
    cells[0].merge(cells[1])
    cells[0].text = "{{ field_of_study }}"
    cells[2].merge(cells[3])
    cells[2].text = "{{ level }}"
    cells[4].text = "{{ profile }}"
    cells[5].merge(cells[6])
    cells[5].text = "{{ semester }}"

    # R10: Zakres
    add_full_row("W zakresie / Specjalizacja magisterska / Moduł kształcenia")
    add_full_row("{{ zakres }}", bold=False)

    # R11: Section Header: RODZAJE ZAJĘĆ
    cell = add_full_row("RODZAJE ZAJĘĆ I ICH WYMIAR GODZINOWY\n(zajęcia dydaktyczne i praca własna studenta)")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Hour rows (simplified for now, matching template structure)
    # Reference image has split columns for stacjonarne/niestacjonarne
    # We'll just put the tag placeholders or labels
    add_full_row("Forma studiów: stacjonarne / niestacjonarne", italic=True)
    # The user didn't mention adding specific hour tags yet, so keeping it similar to current
    add_full_row("... (godziny zajęć) ...", bold=False, italic=True)

    # CEL PRZEDMIOTU
    cell = add_full_row("CEL PRZEDMIOTU*")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_full_row("{{ cel_przedmiotu }}", bold=False)

    # METODY DYDAKTYCZNE
    cell = add_full_row("METODY DYDAKTYCZNE")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_full_row("{{ metody_dydaktyczne }}", bold=False)

    # EFEKTY UCZENIA SIĘ Table Header
    row = table.add_row()
    c0 = row.cells[0]
    for i in range(1, 5): c0.merge(row.cells[i])
    c0.text = "ZAKŁADANE EFEKTY UCZENIA SIĘ PRZEDMIOTU"
    c0.paragraphs[0].runs[0].bold = True
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "Odniesienie do kierunkowych efektów uczenia się"
    c5.paragraphs[0].runs[0].bold = True
    c5.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Wiedza
    row = table.add_row()
    row.cells[0].text = "Wiedza"
    row.cells[0].paragraphs[0].runs[0].bold = True
    c1 = row.cells[1]
    for i in range(2, 5): c1.merge(row.cells[i])
    c1.text = "{{ wiedza }}"
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "{{ learning_outcomesW }}"

    # Umiejętności
    row = table.add_row()
    row.cells[0].text = "Umiejętności"
    row.cells[0].paragraphs[0].runs[0].bold = True
    c1 = row.cells[1]
    for i in range(2, 5): c1.merge(row.cells[i])
    c1.text = "{{ umiejetnosci }}"
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "{{ learning_outcomesU }}"

    # Kompetencje
    row = table.add_row()
    row.cells[0].text = "Kompetencje społeczne"
    row.cells[0].paragraphs[0].runs[0].bold = True
    c1 = row.cells[1]
    for i in range(2, 5): c1.merge(row.cells[i])
    c1.text = "{{ kompetencje }}" # NO "E6" HARDCODED!
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "{{ learning_outcomesK }}"

    # Metody weryfikacji
    row = table.add_row()
    c0 = row.cells[0]
    for i in range(1, 5): c0.merge(row.cells[i])
    c0.text = "Metody weryfikacji efektów uczenia się:\n{{ metody_weryfikacji }}"
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "Symbole efektów przedmiotowych"
    c5.paragraphs[0].runs[0].bold = True

    # TREŚCI KSZTAŁCENIA
    add_full_row("TREŚCI KSZTAŁCENIA")
    add_full_row("{{ content }}", bold=False)

    # Formy zaliczenia
    row = table.add_row()
    c0 = row.cells[0]
    for i in range(1, 5): c0.merge(row.cells[i])
    c0.text = "Formy zaliczenia:\n{{ formy_zaliczenia }}"
    c5 = row.cells[5]
    c5.merge(row.cells[6])
    c5.text = "Procentowy udział w końcowej ocenie"
    c5.paragraphs[0].runs[0].bold = True

    # LITERATURA
    add_full_row("LITERATURA")
    add_full_row("{{ literatura }}", bold=False)

    # Save
    template_path = os.path.join(os.path.dirname(__file__), '..', 'template.docx')
    doc.save(template_path)
    print(f"Rebuilt template saved to {template_path}")

if __name__ == "__main__":
    rebuild()
