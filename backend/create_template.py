from docx import Document

document = Document()

document.add_heading('Sylabus', 0)

document.add_paragraph('Nazwa przedmiotu: {{ name }}')
document.add_paragraph('ECTS: {{ ects }}')
document.add_paragraph('Godziny: {{ hours }}')

document.add_heading('Efekty uczenia się', level=1)
document.add_paragraph('{{ learning_outcomes }}')

document.save('backend/template.docx')
